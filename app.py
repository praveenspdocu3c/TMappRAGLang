from fileinput import filename
import time
import streamlit as st
import json
import pandas as pd
import fitz  # PyMuPDF
from pydantic import BaseModel, Field, ValidationError
from typing import List, Dict, Union
import base64
from docx import Document  
from docx.shared import Pt
from io import BytesIO
import re, ast
import logging
import sys


def parse_trademark_details(document_path: str,) -> List[Dict[str, Union[str, List[int]]]]:
    proposed_class = 30
    trademark_list = []
    
    
    print("+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++ IN FUNCTION ++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++")
    from langchain_community.document_loaders import PyMuPDFLoader
    loader = PyMuPDFLoader(document_path)
    data = loader.load()

    from langchain_text_splitters import RecursiveCharacterTextSplitter
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1024, chunk_overlap=128)
    texts = text_splitter.split_documents(data)
    print("+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++ CHUNK COMPLETED ++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++")
    
    
    from langchain_openai import AzureOpenAIEmbeddings
    from langchain_chroma import Chroma
    embeddings_model = AzureOpenAIEmbeddings(
        model="text-embedding-3-large",
        deployment="text-embedding-3-large",
        api_version="2023-12-01-preview",
        azure_endpoint="https://chat-gpt-a1.openai.azure.com/",
        openai_api_key="c09f91126e51468d88f57cb83a63ee36"
    )    

    print("+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++ Database ++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++")
    # Create a retriever
    retriever = Chroma.from_documents(embedding=embeddings_model, documents=texts).as_retriever(
        search_kwargs={"k": 10}
    )
    print("+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++ RETRIEVER ++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++")
    
    
    
    import fitz  # PyMuPDF
    import json

    def split_text(text, max_length=1500):
        # Split text into chunks of max_length
        chunks = []
        while len(text) > max_length:
            split_point = text.rfind("\n", 0, max_length)
            if split_point == -1:
                split_point = max_length
            chunks.append(text[:split_point])
            text = text[split_point:]
        chunks.append(text)
        return chunks

    from openai import AzureOpenAI
    client = AzureOpenAI(  
                    azure_endpoint="https://chat-gpt-a1.openai.azure.com/",  
                    api_key="c09f91126e51468d88f57cb83a63ee36",  
                    api_version="2024-02-15-preview",
                )  

    def extract_trademark_info_using_llm(text):
        chunks = split_text(text)
        results = []

        for chunk in chunks:
            messages = [
                {
                    "role": "system", "content": """Extract the following details from the text in a structured format: \
                                                - S.no \
                                                - Trademark Name \
                                                - Status \
                                                - International Class number \
                                                - Owner \
                                                - Serial Number \
                                                - Page  
                                                                                            
                                                Note : If the International Class number is not presented instead it has 'Multi' , Then print '['Multi']' as International Class number.
                                                
                                                Note : Only extract if the data is in the same pattern as above mentioned, In Somecases the -Owner details can be splited into two lines, 
                                                If 1st Four trademark details are matched the pattern, Try check 5th line is (integer with '-' & ',') if not then merge the 4th and 5th line as the Owner detail,
                                                Then check for the 6th line as Serial Number and add it
                                                
                                                Print Output in JSON format Example  : [{
                                                                                        "s_no": 1,
                                                                                        "trademark_name": "JOURNEYS",
                                                                                        "status": "REGISTERED",
                                                                                        "international_class_number": [ 25, 35 ],
                                                                                        "owner": "GENSCO BRANDS",
                                                                                        "serial_number": 76-644,691,
                                                                                        "page_number": 15,
                                                                                        },...] """
                },
                {
                    "role": "user", "content": f"Text to extract from:\n{chunk}"
                }
            ]

            response = client.chat.completions.create(  
                    model="DanielChatGPT16k",  
                    messages=messages,  
                    temperature=0,  
                    max_tokens=2000,  
                )  

            results.append(response.choices[0].message.content)

        return results
        

    def read_pdf(file_path: str) -> str:
        document_text = ""
        with fitz.open(file_path) as pdf_document:
            for page_num in range(11, 12):
                page = pdf_document.load_page(page_num)
                page_text = page.get_text()
                document_text += page_text
        return document_text

    # Provide the file path to your PDF document
    pdf_file_path = document_path

    # Extract and parse the details from the PDF
    text = read_pdf(pdf_file_path)
    raw_output = extract_trademark_info_using_llm(text)

    # Store results in the desired format
    Index = []
    for result in raw_output:
        try:
            json_result = json.loads(result)  # Convert string to list of dictionaries
            for item in json_result:
                Index.append({
                    "s_no": item["s_no"],
                    "trademark_name": item["trademark_name"],
                    "status": item["status"],
                    "international_class_number": item["international_class_number"],
                    "owner": item["owner"],
                    "serial_number": item["serial_number"],
                    "page_number": item["page_number"]
                })
        except json.JSONDecodeError as e:
            st.write(f"Error decoding JSON: {result}\nError: {e}")
        except KeyError as e:
            st.write(f"Missing key in result: {result}\nError: {e}")

    print("-----------------------------------------------------------------------------------------------------------------------------------------")
    i = 1
    for item in Index:
        print(f"{i}. Trademark Name: {item['trademark_name']}")
        i = i + 1
    print("-----------------------------------------------------------------------------------------------------------------------------------------")
    
    for item in Index:
        if (item['international_class_number'] == "['Multi']" ):

            trademark_name = {item['trademark_name']}
            t_status = {item['status']}
            t_owner = {item['owner']}
            t_serial_number = {item['serial_number']}

            query = f"""Please extract all International Class numbers for the following trademark details:
                        
                        Trademark Name: {trademark_name}
                        Status: {t_status}
                        Owner: {t_owner}
                        Serial Number: {t_serial_number}
                        
                        Ensure that the entire goods/services section is checked until the phrase "Last Reported Owner:" is encountered. The goods/services may span multiple International class numbers, so please extract all International Class numbers associated with the mentioned trademark.
                        """
            # Get relevant documents ordered by relevance score
            docs = retriever.invoke(query)

            from langchain_community.document_transformers import LongContextReorder
            # Reorder the documents:
            # Less relevant document will be at the middle of the list and more
            # relevant elements at beginning / end.
            reordering = LongContextReorder()
            reordered_docs = reordering.transform_documents(docs)


            from langchain.chains.combine_documents import create_stuff_documents_chain
            from langchain_core.prompts import PromptTemplate
            from langchain_openai import AzureChatOpenAI

            llm = AzureChatOpenAI(
                openai_api_key= "a5c4e09a50dd4e13a69e7ef19d07b48c",
                api_version="2024-02-01",
                azure_endpoint = "https://danielingitaraj.openai.azure.com/",
                model="gpt-4o",
                base_url = None,
                azure_deployment = "GPT4"
            )

            llm.validate_base_url = False

            prompt_template = """
            Given these texts:
            -----
            {context}
            -----
            Please answer the following question:
            {query}
            """
            prompt = PromptTemplate(
                template=prompt_template,
                input_variables=["context", "query"],
            )
            # Create and invoke the chain:
            chain = create_stuff_documents_chain(llm, prompt)
            Class_response = chain.invoke({"context": reordered_docs, "query": query})
            print("++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++")
            print(Class_response)
            print("++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++")


            Flag = "False"

            from openai import AzureOpenAI
            client = AzureOpenAI(  
                            azure_endpoint="https://chat-gpt-a1.openai.azure.com/",  
                            api_key="c09f91126e51468d88f57cb83a63ee36",  
                            api_version="2024-02-15-preview",
                        )  
            messages=[
                            {"role": "system", "content": "You are a helpful assistant for checking the proposed trademark's class number are presented in existing trademark's class number are not."},
                            {"role": "user", "content": "The class number of proposed trademark are: 03. The class number of existing trademark are: The International Class numbers associated with the trademark 'SERENELIFE' are 3, 5, 8, 14, 15, 16, 18, 22, 25, and 35."},
                            {"role": "assistant", "content": "Yes, The International class number of both existing and proposed trademark has 3, Result:True"},
                            {"role": "user", "content": "The class number of proposed trademark are: 45. The class number of existing trademark are: The International Class numbers associated with the trademark 'SERENELIFE' are 3, 5, 8, 14, 15, 16, 18, 22, 25, and 35."},
                            {"role": "assistant", "content": "No, The International class number of both existing and proposed trademark does not has 3, Result:False"},
                            {"role": "user", "content": "The class numbers of the proposed trademark are 18 and 35. The class numbers of the existing trademark 'SOLUTIONS FOR YOUR JOURNEY' are 6, 7, 9, 10, 11, 12, 17, 18, 19, 20, 24, 25, 28, 35, 37, 39, 41, 42, 43, and 44."},
                            {"role": "assistant", "content": "Yes, The International class number of both existing and proposed trademark has 18 and 35, Result:True"},
                            {"role": "user", "content": "The class numbers of the proposed trademark are 18 and 05. The class numbers of the existing trademark 'SOLUTIONS FOR YOUR JOURNEY' are 6, 7, 9, 10, 11, 12, 17, 18, 19, 20, 24, 25, 28, 35, 37, 39, 41, 42, 43, and 44."},
                            {"role": "assistant", "content": "Yes, The International class number of both existing and proposed trademark has 18, Result:True"},
                            {"role": "user", "content": "The class numbers of the proposed trademark are 18 and 35. The class numbers of the existing trademark 'SOLUTIONS FOR YOUR JOURNEY' are 30"},
                            {"role": "assistant", "content": "No, The International class number of both existing and proposed trademark does not has 30, Result:False"},
                            {"role": "user", "content": f"The class number of proposed trademark are: {proposed_class}. The class number of existing trademark are: {Class_response}"}
                        ]
            
            response = client.chat.completions.create(  
                            model="DanielChatGPT16k",  
                            messages=messages,  
                            temperature=0,  
                            max_tokens=1000,  
                        )  

            com_class = response.choices[0].message.content
            # print(com_class)
            Flag = com_class.split("Result:", 1)[1].strip()
            # print(f"Flag : {Flag}")
                    
            if Flag == "True":
                query = f"""Please extract the goods/services of International Class {proposed_class} for the following trademark details:
                                        
                                        Trademark Name: {trademark_name}
                                        Status: {t_status}
                                        Owner: {t_owner}
                                        Serial Number: {t_serial_number}
                                        
                                        Ensure that the entire goods/services section is checked until the phrase "Last Reported Owner:" is encountered. While the goods/services may span multiple International Class numbers, focus on accurately extracting the goods/services of International Class {proposed_class} associated with the mentioned trademark.
                        """
                # Get relevant documents ordered by relevance score
                docs = retriever.invoke(query)


                from langchain_community.document_transformers import LongContextReorder
                # Reorder the documents:
                # Less relevant document will be at the middle of the list and more
                # relevant elements at beginning / end.
                reordering = LongContextReorder()
                reordered_docs = reordering.transform_documents(docs)


                from langchain.chains.combine_documents import create_stuff_documents_chain
                from langchain_core.prompts import PromptTemplate
                from langchain_openai import AzureChatOpenAI

                llm = AzureChatOpenAI(
                    openai_api_key= "a5c4e09a50dd4e13a69e7ef19d07b48c",
                    api_version="2024-02-01",
                    azure_endpoint = "https://danielingitaraj.openai.azure.com/",
                    model="gpt-4o",
                    base_url = None,
                    azure_deployment = "GPT4"
                )

                llm.validate_base_url = False

                prompt_template = """
                Given these texts:
                -----
                {context}
                -----
                Please answer the following question:
                {query}
                """
                prompt = PromptTemplate(
                    template=prompt_template,
                    input_variables=["context", "query"],
                )
                # Create and invoke the chain:
                chain = create_stuff_documents_chain(llm, prompt)
                response = chain.invoke({"context": reordered_docs, "query": query})
                print("++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++")
                print(response)
                print("++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++")                    

                trademark_info = {
                                    "trademark_name": trademark_name,
                                    "owner": t_owner,
                                    "status": t_status,
                                    "serial_number": t_serial_number,
                                    "international_class_number": Class_response,
                                    "goods_services": response,
                            }
                trademark_list.append(trademark_info)

            else :
                
                trademark_info = {
                                    "trademark_name": trademark_name,
                                    "owner": t_owner,
                                    "status": t_status,
                                    "serial_number": t_serial_number,
                                    "international_class_number": Class_response,
                                    "goods_services": "Nill",
                            }
                trademark_list.append(trademark_info)
        
        else :
            
            Flag = "False"

            trademark_name = {item['trademark_name']}
            t_status = {item['status']}
            t_owner = {item['owner']}
            t_serial_number = {item['serial_number']}
            t_class_number = tuple(item['international_class_number'])
            
            
            from openai import AzureOpenAI
            client = AzureOpenAI(  
                            azure_endpoint="https://chat-gpt-a1.openai.azure.com/",  
                            api_key="c09f91126e51468d88f57cb83a63ee36",  
                            api_version="2024-02-15-preview",
                        )  
            messages=[
                            {"role": "system", "content": "You are a helpful assistant for checking the proposed trademark's class number are presented in existing trademark's class number are not."},
                            {"role": "user", "content": "The class number of proposed trademark are: 03. The class number of existing trademark are: The International Class numbers associated with the trademark 'SERENELIFE' are 3, 5, 8, 14, 15, 16, 18, 22, 25, and 35."},
                            {"role": "assistant", "content": "Yes, The International class number of both existing and proposed trademark has 3, Result:True"},
                            {"role": "user", "content": "The class number of proposed trademark are: 45. The class number of existing trademark are: The International Class numbers associated with the trademark 'SERENELIFE' are 3, 5, 8, 14, 15, 16, 18, 22, 25, and 35."},
                            {"role": "assistant", "content": "No, The International class number of both existing and proposed trademark does not has 3, Result:False"},
                            {"role": "user", "content": "The class numbers of the proposed trademark are 18 and 35. The class numbers of the existing trademark 'SOLUTIONS FOR YOUR JOURNEY' are 6, 7, 9, 10, 11, 12, 17, 18, 19, 20, 24, 25, 28, 35, 37, 39, 41, 42, 43, and 44."},
                            {"role": "assistant", "content": "Yes, The International class number of both existing and proposed trademark has 18 and 35, Result:True"},
                            {"role": "user", "content": "The class numbers of the proposed trademark are 18 and 05. The class numbers of the existing trademark 'SOLUTIONS FOR YOUR JOURNEY' are 6, 7, 9, 10, 11, 12, 17, 18, 19, 20, 24, 25, 28, 35, 37, 39, 41, 42, 43, and 44."},
                            {"role": "assistant", "content": "Yes, The International class number of both existing and proposed trademark has 18, Result:True"},
                            {"role": "user", "content": "The class numbers of the proposed trademark are 18 and 35. The class numbers of the existing trademark 'SOLUTIONS FOR YOUR JOURNEY' are 30"},
                            {"role": "assistant", "content": "No, The International class number of both existing and proposed trademark does not has 30, Result:False"},
                            {"role": "user", "content": f"The class number of proposed trademark are: {proposed_class}. The class number of existing trademark are: {t_class_number}"}
                        ]
            
            response = client.chat.completions.create(  
                            model="DanielChatGPT16k",  
                            messages=messages,  
                            temperature=0,  
                            max_tokens=1000,  
                        )  

            com_class = response.choices[0].message.content
            # print(com_class)
            Flag = com_class.split("Result:", 1)[1].strip()
            # print(f"Flag : {Flag}")
                    
            if Flag == "True":
                query = f"""Please extract the goods/services of International Class {proposed_class} for the following trademark details:
                                        
                                        Trademark Name: {trademark_name}
                                        Status: {t_status}
                                        Owner: {t_owner}
                                        Serial Number: {t_serial_number}
                                        
                                        Ensure that the entire goods/services section is checked until the phrase "Last Reported Owner:" is encountered. While the goods/services may span multiple International Class numbers, focus on accurately extracting the goods/services of International Class {proposed_class} associated with the mentioned trademark.
                        """
                # Get relevant documents ordered by relevance score
                docs = retriever.invoke(query)


                from langchain_community.document_transformers import LongContextReorder
                # Reorder the documents:
                # Less relevant document will be at the middle of the list and more
                # relevant elements at beginning / end.
                reordering = LongContextReorder()
                reordered_docs = reordering.transform_documents(docs)


                from langchain.chains.combine_documents import create_stuff_documents_chain
                from langchain_core.prompts import PromptTemplate
                from langchain_openai import AzureChatOpenAI

                llm = AzureChatOpenAI(
                    openai_api_key= "a5c4e09a50dd4e13a69e7ef19d07b48c",
                    api_version="2024-02-01",
                    azure_endpoint = "https://danielingitaraj.openai.azure.com/",
                    model="gpt-4o",
                    base_url = None,
                    azure_deployment = "GPT4"
                )

                llm.validate_base_url = False

                prompt_template = """
                Given these texts:
                -----
                {context}
                -----
                Please answer the following question:
                {query}
                """
                prompt = PromptTemplate(
                    template=prompt_template,
                    input_variables=["context", "query"],
                )
                # Create and invoke the chain:
                chain = create_stuff_documents_chain(llm, prompt)
                response = chain.invoke({"context": reordered_docs, "query": query})
                print("++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++")
                print(response)
                print("++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++")
                
                trademark_info = {
                                    "trademark_name": trademark_name,
                                    "owner": t_owner,
                                    "status": t_status,
                                    "serial_number": t_serial_number,
                                    "international_class_number": t_class_number,
                                    "goods_services": response                                    
                            }
                trademark_list.append(trademark_info)

            else :

                trademark_info = {
                                    "trademark_name": trademark_name,
                                    "owner": t_owner,
                                    "status": t_status,
                                    "serial_number": t_serial_number,
                                    "international_class_number": t_class_number,
                                    "goods_services": "Nill" 
                            }
                trademark_list.append(trademark_info)
                

    return trademark_list
    # print("-----------------------------------------------------------------------------------------------------------------------------------------")
    

def preprocess_text(text: str) -> str:
    text = re.sub(r'\s+', ' ', text).strip()
    text = re.sub(r'[\u2013\u2014]', '-', text)
    return text

def compare_trademarks(existing_trademark: List[Dict[str, Union[str, List[int]]]], proposed_name: str, proposed_class: str, proposed_goods_services: str) -> List[Dict[str, int]]:
    proposed_classes = [int(c.strip()) for c in proposed_class.split(',')]
    messages=[
            {"role": "system", "content": """You are a trademark attorney to properly reasoning based on given conditions and assign conflict grade high or moderate or low to existing trademark and respond with only High or Moderate or Low. \n\n 
                                            Conditions for determining Conflict Grades:\n\n 
                                            
                                            Condition 1: Trademark Name Comparison\n 
                                            - Condition 1A: The existing trademark's name is a character-for-character match with the proposed trademark name.\n 
                                            - Condition 1B: The existing trademark's name is semantically equivalent to the proposed trademark name.\n 
                                            - Condition 1C: The existing trademark's name is phonetically equivalent to the proposed trademark name.\n 
                                            - Condition 1D: Primary Position Requirement- In the context of trademark conflicts, the primary position of a trademark refers to the first word or phrase element in a multi-word or phrase trademark. For a conflict to exist between an existing trademark and a proposed trademark based on Condition 1D, the proposed trademark name must be in the primary position of the existing trademark. This means that the proposed trademark name should be the first word of the existing trademark.\n
                                                            Example:\n Existing Trademark: "STORIES AND JOURNEYS"\n Proposed Trademark: "JOURNEY"\n Analysis:\n The existing trademark "STORIES AND JOURNEYS" consists of multiple words/phrases.\n For the proposed trademark "JOURNEY" to be in conflict under Condition 1D, it must appear as the first word/phrase in the existing trademark.\n In this case, the first word/phrase in "STORIES AND JOURNEYS" is "STORIES", not "JOURNEY".\n Therefore, "JOURNEY" does not meet Condition 1D because it is not in the primary position of the existing trademark.\n
                                                            Example:\n Existing Trademark: "JOURNEY BY COMPANION"\n Proposed Trademark: "JOURNEY"\n Analysis:\n The existing trademark "JOURNEY BY COMPANION" consists of multiple words/phrases.\n For the proposed trademark "JOURNEY" to be in conflict under Condition 1D, it must appear as the first word/phrase in the existing trademark.\n In this case, the first word/phrase in "JOURNEY BY COMPANION" is "JOURNEY".\n Therefore, "JOURNEY" does meet Condition 1D because it is in the primary position of the existing trademark.\n
                                                                                                        
                                            Condition 2: Goods/Services Classification\n 
                                            - Condition 2: The existing trademark's goods/services are in the same class as those of the proposed trademark.\n
                                            
                                            Condition 3: Target Market and Products\n 
                                            - Condition 3A: The existing trademark's goods/services target the exact same products as the proposed trademark.\n 
                                            - Condition 3B: The existing trademark's goods/services target an exact market as the proposed trademark.\n
                                            
                                            If existing trademark in user given input satisfies:\n\n
                                            - Special case: If existing Trademark Status is Cancelled or Abandoned, they will automatically be considered as conflict grade low but still give the reasoning for the potential conflicts.\n\n
                                            - If the existing trademark satisfies Condition 1A, 1B, or 1C, and also satisfies the revised Condition 1D (when applicable), along with Condition 2, and both Condition 3A and 3B, then the conflict grade should be High.\n
                                            - If the existing trademark satisfies any two of the following: Condition 1A, 1B, or 1C (with the revised Condition 1D being a necessary component for these to be considered satisfied when applicable), Condition 2, Condition 3A and 3B, then the conflict grade should be Moderate.\n
                                            - If the existing trademark satisfies only one (or none) of the conditions: Condition 1A, 1B, 1C (only if the revised Condition 1D is also satisfied when applicable), Condition 2, Condition 3A and 3B, then the conflict grade should be Low.\n\n
                                            
                                            Format of the Response:\n
                                            Resoning for Conflict: Reasoning for conflict in bullet points (In reasoning, if exact same goods, services and industries: list the overlaps, you should reasoning whether the good/services are overlapping or not including classes (if same as proposed trademark or not) and see trademark name whether identical (character-for-character) matches, phonetic equivalents, if it is in primary position (first word in the phrase) or not, if it is not in primary position (first word in the phrase) of the existing trademark it is not conflicting and standard plural forms for subject goods and goods that may be related or not. Reasoning should be based on provided information. Do not provide any kind of hypothetical reasoning.)\n\n
                                            
                                            Step 1: Identifying Potential Conflicts
                                            - What is the existing trademark?
                                            - What is the status of the existing trademark?
                                            - Who is the owner of the existing trademark?
                                            - What is the class number for the existing trademark?
                                            - What is the proposed trademark?
                                            - Who is the applicant for the proposed trademark?
                                            - What is the class number for the proposed trademark?

                                            Step 2: Condition 1A - Character-for-Character Match
                                            - Does the existing trademark match the proposed trademark character-for-character?
                                            - If not, does the existing trademark form part of the proposed trademark? Specify the matching part.

                                            Step 3: Condition 1B - Sound-Alike Match
                                            - Do the existing and proposed trademarks sound alike when spoken?
                                            - If yes, describe the similarities in their pronunciation.

                                            Step 4: Condition 1C - Visual Similarity
                                            - Do the existing and proposed trademarks look visually similar?
                                            - If yes, describe the visual elements that contribute to this similarity.

                                            Step 5: Condition 1D - Primary Position
                                            - Is the existing trademark in the primary position of the proposed trademark?
                                            - Is the primary position the beginning of the proposed trademark?
                                            - Does the proposed trademark exactly match the existing trademark in its primary position?
                                            - If the proposed trademark name is a single word or phrase, it must be fully presented in the primary position of the existing trademark for this condition to be applicable.
                                            - As if the existing trademark name is a single word, then condition 1D is not applicable.

                                            Step 6: Condition 2 - Class Overlap
                                            - Do the class numbers of the existing and proposed trademarks match?
                                            - If the class numbers are different, are the classes related in a way that might cause consumer confusion?

                                            Step 7: Condition 3A - Goods/Services Overlap
                                            - What goods/services and products are covered by the existing trademark?
                                            - What goods/services and products are covered by the proposed trademark?
                                            - Is there an exact match or exact overlap between the goods/services and products of the existing and proposed trademarks?

                                            Step 8: Condition 3B - Target Market
                                            - Who is the target market for the goods/services covered by the existing trademark?
                                            - Who is the target market for the goods/services covered by the proposed trademark?
                                            - Is there an exact match or exact overlap in the target market for the existing and proposed trademarks?

                                            Step 9: Conflict Grade Assessment
                                            - Based on the analysis, how would you grade the potential conflict?
                                            - Provide a brief reasoning for the conflict grade.
                                            
                                            Example Analysis Using the Steps :
                                            - Trademark Name: MH
                                            - Trademark Status: REGISTERED
                                            - Trademark Owner: ZHAO
                                            - Trademark Class Number: 3
                                            - Proposed Trademark: MH BY MOTHERHOOD
                                            - Applicant: ABC Company
                                            - Proposed Trademark Class Number: 3

                                            Step 2: Condition 1A - Character-for-Character Match
                                            - Does the existing trademark match the proposed trademark character-for-character?
                                            - No, "MH" is part of "MH BY MOTHERHOOD" but not an exact match.

                                            Step 3: Condition 1B - Sound-Alike Match
                                            - Do the existing and proposed trademarks sound alike when spoken?
                                            - Yes, The existing trademark "MH" and the proposed trademark "MH BY MOTHERHOOD" are phonetically equivalent
                                            - If the existing trademark name is multi-word, then the primary word of the existing trademark name followed by a color or the name of a country, city, or number (either numeric or letter) should be considered as Condition satisfied.

                                            Step 4: Condition 1C - Visual Similarity
                                            - Do the existing and proposed trademarks look visually similar?
                                            - No, "MH" and "MH BY MOTHERHOOD" do not look visually similar.

                                            Step 5: Condition 1D - Primary Position
                                            - Is the existing trademark in the primary position of the proposed trademark?
                                            - Yes, the proposed trademark "MH BY MOTHERHOOD", which has "MH" as in the primary position of the existing trademark "MH".

                                            Step 6: Condition 2 - Class Overlap
                                            - Do the class numbers of the existing and proposed trademarks match?
                                            - Yes, both are in Class 3.

                                            Step 7: Condition 3A - Goods/Services Overlap
                                            - What goods/services are covered by the existing trademark?
                                            - Non-medicated cosmetic and hair care preparations.
                                            - What goods/services are covered by the proposed trademark?
                                            - Hair care preparations.
                                            - Is there an overlap between the goods/services and product of the existing and proposed trademarks?
                                            - Yes, both include exact products such as hair care preparations.

                                            Step 8: Condition 3B - Target Market
                                            - Who is the target market for the goods/services covered by the existing trademark?
                                            - Consumers interested in cosmetic and hair care products.
                                            - Who is the target market for the goods/services covered by the proposed trademark?
                                            - Consumers interested in hair care products.
                                            - Is there an overlap in the target market for the existing and proposed trademarks?
                                            - Yes, both target the exact same consumers interested in hair care products.

                                            Step 9: Conflict Grade Assessment
                                            - Based on the analysis, how would you grade the potential conflict?
                                            - High.

                                            Provide a brief reasoning for the conflict grade.
                                            - Example : Condition 1A is not satisfied since there is no character-for-character match and 1C are not satisfied either. However, Conditions 1B and Condition 1D is been satisfied and also there is a class overlap and goods/services overlap, as well as a shared target market, indicating a High potential for consumer confusion.
                                                                                    
                                            Conflict Grade: Based on above reasoning (Low or Moderate or High)."""
                                            },
            
            {"role": "user", "content": """Compare the following existing and proposed trademarks and determine the conflict grade.\n
                                            Existing Trademark:\n
                                            Name: DISCOVER WHAT’S NEXT\n
                                            Goods/Services: Online retail store services geared for men and women, featuring a wide variety of unique consumer products\n 
                                            International Class Numbers: 35\n
                                            Status: REGISTERED\n
                                            Owner: THE GIDDYUP GROUP, INC\n
                                            
                                            Proposed Trademark:\n
                                            Name: DISCOVER\n 
                                            Goods/Services: Luggage and carrying bags; suitcases, trunks, travelling bags, sling bags for = carrying infants, school bags; purses; wallets; retail and online retail services\n
                                            International Class Numbers: 18, 35\n"""
            },
        {"role": "assistant", "content":"""
Reasoning for Conflict:
Step 1: Condition 1A - Exact Character Match
- Does the existing trademark contain the proposed trademark term?
- Yes, the existing trademark "DISCOVER WHAT’S NEXT" contains the term "DISCOVER," which is a character for-character match with the proposed trademark "DISCOVER."
- Condition 1A is satisfied.

Step 2: Condition 1D - Primary Position
- Is the existing trademark in the primary position of the proposed trademark?
- Yes, "DISCOVER" is in the primary position of the existing trademark "DISCOVER WHAT’S NEXT."
- As if the existing trademark name is a single word, then condition 1D is not applicable.
- Existing trademark name is a single word, Hence Condition 1D is not applicable.

Step 3: Condition 2 - Class Overlap
- Do the class numbers of the existing and proposed trademarks overlap?
- Yes, the existing trademark is registered under International Class 35, which overlaps with the proposed trademark's Class 35 for retail and online retail services.
- Condition 2 is satisfied.

Step 4: Condition 3A - Goods/Services Overlap
- What goods/services and products are covered by the existing trademark?
- The existing trademark's goods/services include online retail store services geared for men and women, featuring a wide variety of unique consumer products.
- What goods/services and products are covered by the proposed trademark?
- The proposed trademark's goods/services include luggage and carrying bags; suitcases, trunks, travelling bags, sling bags for carrying infants, school bags; purses; wallets; retail and online retail services.
- Is there an exact match or exact overlap between the goods/services and products of the existing and proposed trademarks?
- No, while both trademarks include retail and online retail services, the existing trademark focuses on a wide variety of unique consumer products, whereas the proposed trademark focuses specifically on luggage, bags, and related items.
- Condition 3A is not fully satisfied.

Step 5: Condition 3B - Target Market Overlap
- What is the target market for the existing trademark?
- The existing trademark targets consumers interested in a wide variety of unique consumer products.
- What is the target market for the proposed trademark?
- The proposed trademark targets consumers interested in luggage, carrying bags, and related products.
- Is there an exact overlap in the target market for the existing and proposed trademarks?
- No, although there is some overlap in the general market of retail and online retail services, the specific focus of the products differs.
- Condition 3B is not fully satisfied.

Conclusion:
- Despite satisfying Conditions 1A, (1D not applicable) , and 2, the differences in the specific goods/services and target markets mean that the conflict is not high.
- Since Conditions 3A and 3B are not fully satisfied due to the lack of exact overlap in the goods/services and target markets, the conflict grade should be moderate rather than high.

Reason:
Reasoning for Conflict:
1A - The existing trademark "DISCOVER WHAT’S NEXT" contains the term "DISCOVER," which is a character for-character match with the proposed trademark "DISCOVER."
1D - "DISCOVER" is in the primary position of the existing trademark "DISCOVER WHAT’S NEXT". However, - Existing trademark name is a single word, Hence Condition 1D is not applicable. 
2  - The existing trademark is registered under International Class 35, which overlaps with the proposed trademark's Class 35 for retail and online retail services.
3A - Both trademarks include retail and online retail services, the existing trademark focuses on a wide variety of unique consumer products, whereas the proposed trademark focuses specifically on luggage, bags, and related items.
3B - Although there is some overlap in the general market of retail and online retail services, the specific focus of the products differs.

Conclusion:
- Despite satisfying Conditions 1A, (1D not applicable) , and 2, the differences in the specific goods/services and target markets mean that the conflict is not high.
- Since Conditions 3A and 3B are not fully satisfied due to the lack of exact overlap in the goods/services and target markets, the conflict grade should be moderate rather than high.

- Conflict Grade: Moderate

"""
            },
            {"role": "user", "content": """Compare the following existing and proposed trademarks and determine the conflict grade.\n
                                            Existing Trademark:\n
                                            Name: DB JOURNEY\n
                                            Goods/Services: All-purpose carrying bags, rucksacks, hipsacks, luggage, toiletry bags, key bags, luggage tags, pocket wallets, straps for luggage, shoulder straps, and umbrellas (Class 18)\n 
                                            International Class Numbers: 9, 16, 18, 25, 28\n
                                            Status: PENDING SECTION 66(A) (MADRID PROTOCOL)\n
                                            Owner: DB EQUIPMENT AS NORWAY AS\n
                                            
                                            Proposed Trademark:\n
                                            Name: JOURNEY\n 
                                            Goods/Services: Luggage and carrying bags; suitcases, trunks, travelling bags, sling bags for carrying infants, school bags; purses; wallets; retail and online retail services\n
                                            International Class Numbers: 18, 35\n"""
            },
            {"role": "assistant", "content":""" 
Reasoning for Conflict:
Step 1: Condition 1A - Character-for-Character Match
- The existing trademark "DB JOURNEY" is not a character-for-character match with the proposed trademark "JOURNEY."
- Condition 1A is not satisfied.

Step 2: Condition 1B - Semantic Equivalence
- The existing trademark "DB JOURNEY" and the proposed trademark "JOURNEY" are not semantically equivalent.
- Condition 1B is not satisfied.

Step 3: Condition 1C - Phonetic Equivalence
- The existing trademark "DB JOURNEY" and the proposed trademark "JOURNEY" are phonetically similar due to the shared term "JOURNEY."
- Condition 1C is satisfied.

Step 4: Condition 1D - Primary Position
- The term "JOURNEY" is in the primary position in the proposed trademark "JOURNEY."
- The term "JOURNEY" is not in the primary position in the existing trademark "DB JOURNEY" (the primary term is "DB").
- Condition 1D is not satisfied.

Step 5: Condition 2 - Class Overlap
- The existing trademark includes Class 18, which overlaps with the proposed trademark's Class 18 for luggage and carrying bags.
- The existing trademark does not include Class 35, which is part of the proposed trademark's classification.
- Condition 2 is partially satisfied.

Step 6: Condition 3A - Goods/Services Overlap
- What goods/services and products are covered by the existing trademark?
- The existing trademark's goods/services in Class 18 include all-purpose carrying bags, rucksacks, hipsacks, luggage, toiletry bags, key bags, luggage tags, pocket wallets, straps for luggage, shoulder straps, and umbrellas.
- What goods/services and products are covered by the proposed trademark?
- The proposed trademark's goods/services in Class 18 include luggage and carrying bags, suitcases, trunks, travelling bags, sling bags for carrying infants, school bags, purses, and wallets.
- Is there an exact match or exact overlap between the goods/services and products of the existing and proposed trademarks?
- There is a significant overlap in the goods/services in Class 18.
- Condition 3A is satisfied.

Step 7: Condition 3B - Target Market Overlap
- What is the target market for the existing trademark?
- The existing trademark targets consumers interested in a wide range of bags and carrying cases, including luggage and related accessories.
- What is the target market for the proposed trademark?
- The proposed trademark targets consumers interested in luggage, carrying bags, and related products.
- Is there an exact overlap in the target market for the existing and proposed trademarks?
- There is an overlap in the target market for the existing and proposed trademarks.
- Condition 3B is satisfied.

Conclusion:
- Conditions 1C, 2 (partially), 3A, and 3B are satisfied.
- Condition 1D is not satisfied due to the proposed trademark name "JOURNEY" not being in the primary position of the existing trademark name "DB JOURNEY."
- Given the significant overlap in goods/services and the shared target market but the non-satisfaction of - Condition 1D, the conflict grade should be moderate rather than high.

Reason:
Reasoning for Conflict:
1A - The existing trademark "DB JOURNEY" is not a character-for-character match with the proposed trademark "JOURNEY."
1B - The existing trademark "DB JOURNEY" and the proposed trademark "JOURNEY" are not semantically equivalent.
1C - The existing trademark "DB JOURNEY" and the proposed trademark "JOURNEY" are phonetically similar due to the shared term "JOURNEY."
1D - The term "JOURNEY" is not in the primary position in the existing trademark "DB JOURNEY" (the primary term is "DB").
2  - The existing trademark includes Class 18, which overlaps with the proposed trademark's Class 18 for luggage and carrying bags.
3A - There is a significant overlap in the goods/services in Class 18.
3B - There is an overlap in the target market for the existing and proposed trademarks.

Conclusion:
- Conditions 1C, 2 (partially), 3A, and 3B are satisfied.
- Condition 1D is not satisfied due to the proposed trademark name "JOURNEY" not being in the primary position of the existing trademark name "DB JOURNEY."
- Given the significant overlap in goods/services and the shared target market but the non-satisfaction of - Condition 1D, the conflict grade should be moderate rather than high.

- Conflict Grade: Moderate

"""
            },
            {"role": "user", "content": """Compare the following existing and proposed trademarks and determine the conflict grade.\n
                                            Existing Trademark:\n
                                            Name: RADIANT RED\n
                                            Goods/Services: Non−medicated hair care preparations\n 
                                            International Class Numbers: 3\n
                                            Status: REGISTERED\n
                                            Owner: KAO KABUSHIKI KAISHA TA KAO CORPORATION JAPAN CORPORATION\n
                                            
                                            Proposed Trademark:\n
                                            Name: RADIANT AMBER\n 
                                            Goods/Services: DEODORANT\n
                                            International Class Numbers: 3\n"""
            },
            {"role": "assistant", "content":"""  
Reasoning for Conflict:
Step 1: Condition 1A - Character-for-Character Match
- The existing trademark "RADIANT RED" is not a character-for-character match with the proposed trademark "RADIANT AMBER."
- Condition 1A is not satisfied.

Step 2: Condition 1B - Semantic Equivalence
- The existing trademark "RADIANT RED" and the proposed trademark "RADIANT AMBER" are not semantically equivalent.
- Condition 1B is not satisfied.

Step 3: Condition 1C - Phonetic Equivalence
- The existing trademark "RADIANT RED" and the proposed trademark "RADIANT AMBER" are phonetically similar due to the shared term "RADIANT."
- Condition 1C is satisfied.

Step 4: Condition 1D - Primary Position
- The term "RADIANT" is in the primary position in the proposed trademark "RADIANT AMBER."
- The term "RADIANT" is in the primary position in the existing trademark "RADIANT RED."
- Both the existing trademark and the proposed trademark have "RADIANT" in their primary positions.
- If the proposed trademark name is a single word or phrase, it must be fully presented in the primary position of the existing trademark for this condition to be applicable.
- If the existing trademark name is a single word, then condition 1D is not applicable.
- Condition 1D is satisfied.

Step 5: Condition 2 - Class Overlap
- Both the existing and proposed trademarks are in International Class 3.
- Condition 2 is satisfied.

Step 6: Condition 3A - Goods/Services Overlap
- What goods/services and products are covered by the existing trademark?
- The existing trademark's goods/services include non-medicated hair care preparations.
- What goods/services and products are covered by the proposed trademark?
- The proposed trademark's goods/services include deodorant.
- Is there an exact match or exact overlap between the goods/services and products of the existing and proposed trademarks?
- No, there is no exact match or overlap as the existing trademark covers hair care preparations and the proposed trademark covers deodorant.
- Condition 3A is not satisfied.

Step 7: Condition 3B - Target Market Overlap
- What is the target market for the existing trademark?
- The existing trademark targets consumers interested in non-medicated hair care preparations.
- What is the target market for the proposed trademark?
- The proposed trademark targets consumers interested in deodorant.
- Is there an exact overlap in the target market for the existing and proposed trademarks?
- No, there is no exact overlap in the target market as they focus on different types of personal care products.
- Condition 3B is not satisfied.

Conclusion:
- Conditions 1C, 1D, and 2 are satisfied.
- Conditions 1A, 1B, 3A, and 3B are not satisfied.
- Given the phonetic similarity, primary position match, and class overlap but the significant differences in goods/services and target markets, the conflict grade should be moderate.

Reason:
1A - The existing trademark "RADIANT RED" is not a character-for-character match with the proposed trademark "RADIANT AMBER."
1B - The existing trademark "RADIANT RED" and the proposed trademark "RADIANT AMBER" are not semantically equivalent.
1C - The existing trademark "RADIANT RED" and the proposed trademark "RADIANT AMBER" are phonetically similar due to the shared term "RADIANT."
1D - Both the existing trademark and the proposed trademark have "RADIANT" in their primary positions.
2  - Both trademarks are in International Class 3.
3A - There is no exact match or overlap between the goods/services in Class 3.
3B - There is no exact overlap in the target market for the existing and proposed trademarks.

Conclusion:
- Conditions 1C, 1D, and 2 are satisfied.
- Conditions 1A, 1B, 3A, and 3B are not satisfied.
- Given the phonetic similarity, primary position match, and class overlap but the significant differences in goods/services and target markets, the conflict grade should be moderate.

- Conflict Grade: Moderate             

            """},
            {"role": "user", "content": """Compare the following existing and proposed trademarks and determine the conflict grade.\n
                                            Existing Trademark:\n
                                            Name: MH\n
                                            Goods/Services: Non−medicated cosmetic hair care preparations in the nature of hair wax; cosmetic hair filling fibers for covering bald and thinning spots on the scalp\n 
                                            International Class Numbers: 3\n
                                            Status: REGISTERED\n
                                            Owner: ZHAO\n
                                            
                                            Proposed Trademark:\n
                                            Name: MH BY MOTHERHOOD\n 
                                            Goods/Services: IC 003: SKIN CARE PREPARATIONS; COSMETICS; BABY CARE PRODUCTS, NAMELY, SKIN SOAPS, BABY WASH, BABY BUBBLE BATH, BABY LOTIONS, BABY SHAMPOOS; SKIN CLEANSERS; BABY WIPES; NON− MEDICATED DIAPER RASH OINTMENTS AND LOTIONS; SKIN LOTIONS, CREAMS, MOISTURIZERS, AND OILS; BODY WASH; BODY SOAP; DEODORANTS; PERFUME; HAIR CARE PREPARATIONS\n
                                            International Class Numbers: 3\n"""
            },
            {"role": "assistant", "content":""" 
Reasoning for Conflict:
Step 1: Condition 1A - Character-for-Character Match
- The existing trademark "MH" is not a character-for-character match with the proposed trademark "MH BY MOTHERHOOD."
- Condition 1A is not satisfied.

Step 2: Condition 1B - Semantic Equivalence
- The existing trademark "MH" and the proposed trademark "MH BY MOTHERHOOD" are not semantically equivalent.
- Condition 1B is not satisfied.

Step 3: Condition 1C - Phonetic Equivalence
- The existing trademark "MH" and the proposed trademark "MH BY MOTHERHOOD" are phonetically similar due to the shared term "MH."
- Condition 1C is satisfied.

Step 4: Condition 1D - Primary Position
- The term "MH" is in the primary position in the proposed trademark "MH BY MOTHERHOOD."
- The term "MH" is in the primary position in the existing trademark "MH."
- As if the existing trademark name is a single word, then condition 1D is not applicable.
- Existing trademark name is a single word, Hence Condition 1D is not applicable.
- Condition 1D is not applicable.

Step 5: Condition 2 - Class Overlap
- Both the existing and proposed trademarks are in International Class 3.
- Condition 2 is satisfied.

Step 6: Condition 3A - Goods/Services Overlap
- What goods/services and products are covered by the existing trademark?
- The existing trademark's goods/services include non-medicated cosmetic hair care preparations in the nature of hair wax and cosmetic hair filling fibers for covering bald and thinning spots on the scalp.
- What goods/services and products are covered by the proposed trademark?
- The proposed trademark's goods/services include skin care preparations, cosmetics, baby care products, skin cleansers, baby wipes, non-medicated diaper rash ointments and lotions, skin lotions, creams, moisturizers, oils, body wash, body soap, deodorants, perfume, and hair care preparations.
- Is there an exact match or exact overlap between the goods/services and products of the existing and proposed trademarks?
- There is a partial overlap in the goods/services, specifically in the area of hair care preparations.
- Condition 3A is partially satisfied.

Step 7: Condition 3B - Target Market Overlap
- What is the target market for the existing trademark?
- The existing trademark targets consumers interested in non-medicated cosmetic hair care preparations.
- What is the target market for the proposed trademark?
- The proposed trademark targets consumers interested in a broader range of skin care, cosmetics, baby care products, and hair care preparations.
- Is there an exact overlap in the target market for the existing and proposed trademarks?
- There is a partial overlap in the target market, specifically for consumers interested in hair care preparations.
- Condition 3B is partially satisfied.

Conclusion:
- Conditions 1C, 2, and (1D not applicable) are satisfied.
- Conditions 1A, 1B are not satisfied.
- Conditions 3A and 3B are partially satisfied due to the overlap in hair care preparations and the shared target market for those products.
- Given the partial overlap in goods/services and the shared target market but the lack of satisfaction in key similarity conditions (1A, 1B), the conflict grade should be moderate rather than high.

Reason:
Reasoning for Conflict:
1A - The existing trademark "MH" is not a character-for-character match with the proposed trademark "MH BY MOTHERHOOD."
1B - The existing trademark "MH" and the proposed trademark "MH BY MOTHERHOOD" are not semantically equivalent.
1C - The existing trademark "MH" and the proposed trademark "MH BY MOTHERHOOD" are phonetically similar due to the shared term "MH."
1D - The term "MH" is in the primary position in the proposed trademark "MH BY MOTHERHOOD." and the term "MH" is in the primary position in the existing trademark "MH." However, Existing trademark name is a single word, Hence Condition 1D is not applicable.
2  - Both the existing and proposed trademarks are in International Class 3.
3A - There is a partial overlap in the goods/services, specifically in the area of hair care preparations.
3B - There is a partial overlap in the target market, specifically for consumers interested in hair care preparations.

Conclusion:
- Conditions 1C, 2, and (1D not applicable) are satisfied.
- Conditions 1A, 1B are not satisfied.
- Conditions 3A and 3B are partially satisfied due to the overlap in hair care preparations and the shared target market for those products.
- Given the partial overlap in goods/services and the shared target market but the lack of satisfaction in key similarity conditions (1A, 1B), the conflict grade should be moderate rather than high.

- Conflict Grade: Moderate

"""         },            
            {"role": "user", "content": """Compare the following existing and proposed trademarks and determine the conflict grade.\n
                                            Existing Trademark:\n
                                            Name: SCOOPT'D\n
                                            Goods/Services: Hypoallergenic and vegan-friendly ice cream\n 
                                            International Class Numbers: 30\n
                                            Status: Registered\n
                                            Owner: Scoopt'd: Dairy Free Treats\n
                                            
                                            Proposed Trademark:\n
                                            Name: SCOOP-A-PALOOZA\n 
                                            Goods/Services: Ice cream\n
                                            International Class Numbers: 30\n"""
            },
            {"role": "assistant", "content":""" 
Reasoning for Conflict:
Step 1: Condition 1A - Character-for-Character Match
- The existing trademark "SCOOPT'D" is not a character-for-character match with the proposed trademark "SCOOP-A-PALOOZA."
- Condition 1A is not satisfied.

Step 2: Condition 1B - Semantic Equivalence
- The existing trademark "SCOOPT'D" and the proposed trademark "SCOOP-A-PALOOZA" are not semantically equivalent.
- Condition 1B is not satisfied.

Step 3: Condition 1C - Phonetic Equivalence
- The existing trademark "SCOOPT'D" and the proposed trademark "SCOOP-A-PALOOZA" are partially phonetically similar due to the shared term "SCOOP," which could lead to confusion. However, they are not fully phonetically similar.
- Condition 1C is not satisfied.

Step 4: Condition 1D - Primary Position
- The term "SCOOP" is in the primary position in the proposed trademark "SCOOP-A-PALOOZA."
- The term "SCOOP" is in the primary position in the existing trademark "SCOOPT'D."
- The full proposed trademark "SCOOP-A-PALOOZA" does not match the primary position of the existing trademark "SCOOPT'D."
- As if the existing trademark name is a single word, then condition 1D is not applicable.
- Existing trademark name is a single word, Hence Condition 1D is not applicable.
- Condition 1D is not applicable.

Step 5: Condition 2 - Class Overlap
- Both the existing and proposed trademarks are in International Class 30.
- Condition 2 is satisfied.

Step 6: Condition 3A - Goods/Services Overlap
- What goods/services and products are covered by the existing trademark?
- The existing trademark's goods/services include hypoallergenic and vegan-friendly ice cream.
- What goods/services and products are covered by the proposed trademark?
- The proposed trademark's goods/services include ice cream.
- Is there an exact match or exact overlap between the goods/services and products of the existing and proposed trademarks?
- Yes, both trademarks cover ice cream.
- Condition 3A is satisfied.

Step 7: Condition 3B - Target Market Overlap
- What is the target market for the existing trademark?
- The existing trademark targets consumers interested in hypoallergenic and vegan-friendly ice cream.
- What is the target market for the proposed trademark?
- The proposed trademark targets consumers interested in ice cream.
- Is there an exact overlap in the target market for the existing and proposed trademarks?
- Yes, both trademarks target consumers interested in ice cream.
- Condition 3B is satisfied.

Conclusion:
- Conditions 2, 3A, and 3B and (1D not applicable) are satisfied.
- Conditions 1A, 1B, 1C are not satisfied.
- Given the overlap in goods/services and the shared target market but the lack of satisfaction in key similarity conditions (1A, 1B, 1C, and 1D), the conflict grade should be moderate rather than high.

Reason:
Reasoning for Conflict:
1A - The existing trademark "SCOOPT'D" is not a character-for-character match with the proposed trademark "SCOOP-A-PALOOZA."
1B - The existing trademark "SCOOPT'D" and the proposed trademark "SCOOP-A-PALOOZA" are not semantically equivalent.
1C - The existing trademark "SCOOPT'D" and the proposed trademark "SCOOP-A-PALOOZA" are partially phonetically similar due to the shared term "SCOOP," which could lead to confusion. However, they are not fully phonetically similar.
1D - The full proposed trademark "SCOOP-A-PALOOZA" does not match the primary position of the existing trademark "SCOOPT'D." However, Existing trademark name is a single word, Hence Condition 1D is not applicable.
2  - Both the existing and proposed trademarks are in International Class 30.
3A - Both trademarks cover ice cream.
3B - Both trademarks target consumers interested in ice cream.

Conclusion:
- Conditions 2, 3A, and 3B and (1D not applicable) are satisfied.
- Conditions 1A, 1B, 1C are not satisfied.
- Given the overlap in goods/services and the shared target market but the lack of satisfaction in key similarity conditions (1A, 1B, 1C, and 1D), the conflict grade should be moderate rather than high.

- Conflict Grade: Moderate

"""
            },
            {"role": "user", "content": f"""Compare the following existing and proposed trademarks and determine the conflict grade.\n
                                            Existing Trademark:\n
                                            Name: {existing_trademark['trademark_name']}\n
                                            Goods/Services: {existing_trademark['goods_services']}\n 
                                            International Class Numbers: {existing_trademark['international_class_number']}\n
                                            Status: {existing_trademark['status']}\n
                                            Owner: {existing_trademark['owner']}\n
                                            Proposed Trademark:\n
                                            Name: {proposed_name}\n 
                                            Goods/Services: {proposed_goods_services}\n
                                            International Class Numbers: {proposed_classes}\n"""
            }
        ]
                
    from openai import AzureOpenAI
    client = AzureOpenAI(  
                azure_endpoint="https://danielingitaraj.openai.azure.com/",  
                api_key="a5c4e09a50dd4e13a69e7ef19d07b48c",  
                api_version="2024-02-01",
            )  
                
    response_reasoning = client.chat.completions.create(  
                        model="GPT4",  
                        messages=messages,  
                        temperature=0,  
                        max_tokens=4000,  
                        top_p = 1
                    )  

    Treasoning = response_reasoning.choices[0].message.content
    reasoning = Treasoning.split("Reason:", 1)[1].strip()
    conflict_grade = Treasoning.split("Conflict Grade:", 1)[1].strip() 
    progress_bar.progress(70)

    return {
        'Trademark name': existing_trademark['trademark_name'],
        'Trademark status': existing_trademark['status'],
        'Trademark owner': existing_trademark['owner'],
        'Trademark class Number': existing_trademark['international_class_number'],
        'Trademark serial number' : existing_trademark['serial_number'],
        'conflict_grade': conflict_grade,
        'reasoning': reasoning
    }


def extract_proposed_trademark_details(file_path: str) -> Dict[str, Union[str, List[int]]]:
    """ Extract proposed trademark details from the given input format """
    proposed_details = {}
    with fitz.open(file_path) as pdf_document:
        if pdf_document.page_count > 0:
            page = pdf_document.load_page(0)
            page_text = preprocess_text(page.get_text())
            
    name_match = re.search(r'Mark Searched:\s*(.*?)(?=\s*Client Name:)', page_text, re.IGNORECASE | re.DOTALL)
    if name_match:
        proposed_details["proposed_trademark_name"] = name_match.group(1).strip()

    goods_services_match = re.search(r'Goods/Services:\s*(.*?)(?=\s*Trademark Research Report)', page_text, re.IGNORECASE | re.DOTALL)
    if goods_services_match:
        proposed_details["proposed_goods_services"] = goods_services_match.group(1).strip()
    
    # Use LLM to find the international class number based on goods & services
    if "proposed_goods_services" in proposed_details:
        goods_services = proposed_details["proposed_goods_services"]
        class_numbers = find_class_numbers(goods_services)
        proposed_details["proposed_nice_classes_number"] = class_numbers
    
    return proposed_details

def find_class_numbers(goods_services: str) -> List[int]:
    """ Use LLM to find the international class numbers based on goods & services """
        # Initialize AzureChatOpenAI
    from openai import AzureOpenAI
    client = AzureOpenAI(  
                    azure_endpoint="https://chat-gpt-a1.openai.azure.com/",  
                    api_key="c09f91126e51468d88f57cb83a63ee36",  
                    api_version="2024-02-15-preview",
                )
    messages=[
            {"role": "system", "content": "You are a helpful assistant for finding the International class number of provided Goods & Services."},
            {"role": "user", "content": "The goods/services are: IC 003: SKIN CARE PREPARATIONS; COSMETICS; BABY CARE PRODUCTS, NAMELY, SKIN SOAPS, BABY WASH, BABY BUBBLE BATH, BABY LOTIONS, BABY SHAMPOOS; SKIN CLEANSERS; BABY WIPES; NON− MEDICATED DIAPER RASH OINTMENTS AND LOTIONS; SKIN LOTIONS, CREAMS, MOISTURIZERS, AND OILS; BODY WASH; BODY SOAP; DEODORANTS; PERFUME; HAIR CARE PREPARATIONS. Find the international class numbers."},
            {"role": "assistant", "content": "The international class numbers : 03"},
            {"role": "user", "content": "The goods/services are: LUGGAGE AND CARRYING BAGS; SUITCASES, TRUNKS, TRAVELLING BAGS, SLING BAGS FOR CARRYING INFANTS, SCHOOL BAGS; PURSES; WALLETS; RETAIL AND ONLINE RETAIL SERVICES. Find the international class numbers."},
            {"role": "assistant", "content": "The international class numbers : 18,35"},
            {"role": "user", "content": f"The goods/services are: {goods_services}. Find the international class numbers."}
        ]
    
    response = client.chat.completions.create(  
                        model="DanielChatGPT16k",  
                        messages=messages,  
                        temperature=0,  
                        max_tokens=200,  
    )  

    class_numbers_str = response.choices[0].message.content
    
    # Extracting class numbers and removing duplicates
    class_numbers = re.findall(r'(?<!\d)\d{2}(?!\d)', class_numbers_str)  # Look for two-digit numbers
    class_numbers = ','.join(set(class_numbers))  # Convert to set to remove duplicates, then join into a single string
    
    return class_numbers

def extract_proposed_trademark_details2(file_path: str) -> Dict[str, Union[str, List[int]]]:
    """ Extract proposed trademark details from the first page of the document """
    proposed_details = {}
    with fitz.open(file_path) as pdf_document:
        if pdf_document.page_count > 0:
            page = pdf_document.load_page(0)
            page_text = preprocess_text(page.get_text())
            
            name_match = re.search(r'Name:\s*(.*?)(?=\s*Nice Classes:)', page_text)
            if name_match:
                proposed_details["proposed_trademark_name"] = name_match.group(1).strip()
                
            nice_classes_match = re.search(r'Nice Classes:\s*(\d+(?:,\s*\d+)*)', page_text)
            if nice_classes_match:
                proposed_details["proposed_nice_classes_number"] = nice_classes_match.group(1).strip()
            
            goods_services_match = re.search(r'Goods & Services:\s*(.*?)(?=\s*Registers|$)', page_text, re.IGNORECASE | re.DOTALL)
            if goods_services_match:
                proposed_details["proposed_goods_services"] = goods_services_match.group(1).strip()
    
    return proposed_details

def list_conversion(proposed_class: str) -> List[int]:
    
    from openai import AzureOpenAI
    client = AzureOpenAI(  
                    azure_endpoint="https://chat-gpt-a1.openai.azure.com/",  
                    api_key="c09f91126e51468d88f57cb83a63ee36",  
                    api_version="2024-02-15-preview",
                )

    messages=[      {"role": "system", "content": "You are a helpful assistant for converting the class number string into python list of numbers.\n Respond only with python list. Example : [18,35]"},
                    {"role": "user", "content": "The class number are: 15,89. convert the string into python list of numbers."},
                    {"role": "assistant", "content": "[15,89]"},
                    {"role": "user", "content": f"The class number are: {proposed_class}. convert the string into python list of numbers."}
                ]
    
    response = client.chat.completions.create(  
                        model="DanielChatGPT16k",  
                        messages=messages,  
                        temperature=0,  
                        max_tokens=200,  
                    )  

    lst_class = response.choices[0].message.content
    class_value = ast.literal_eval(lst_class)
            
    return class_value

# Streamlit App  
st.title("Trademark Document Parser RAG")  
  
# File upload  
uploaded_files = st.sidebar.file_uploader("Choose PDF files", type="pdf", accept_multiple_files=True)  
  
if uploaded_files:  
    if st.sidebar.button("Check Conflicts", key="check_conflicts"):  
        total_files = len(uploaded_files)  
        progress_bar = st.progress(0)  
        for i, uploaded_file in enumerate(uploaded_files):  
            # Save uploaded file to a temporary file path  
            temp_file_path = f"temp_{uploaded_file.name}"  
            with open(temp_file_path, "wb") as f:  
                f.write(uploaded_file.read())  
            
            sp = True
            proposed_trademark_details = extract_proposed_trademark_details(temp_file_path)  
                            
            if proposed_trademark_details:  
                proposed_name = proposed_trademark_details.get('proposed_trademark_name', 'N')  
                proposed_class = proposed_trademark_details.get('proposed_nice_classes_number')  
                proposed_goods_services = proposed_trademark_details.get('proposed_goods_services', 'N') 
                if (proposed_goods_services != 'N'): 
                    with st.expander(f"Proposed Trademark Details for {uploaded_file.name}"):  
                            st.write(f"Proposed Trademark name: {proposed_name}")  
                            st.write(f"Proposed class-number: {proposed_class}")  
                            st.write(f"Proposed Goods & Services: {proposed_goods_services}") 
                    class_list = list_conversion(proposed_class) 
                else :
                    st.write("______________________________________________________________________________________________________________________________")
                    st.write(f"Sorry, unable to generate report due to insufficient information about goods & services in the original trademark report : {uploaded_file.name}")
                    st.write("______________________________________________________________________________________________________________________________")
                    sp = False
            else:  
                
                proposed_trademark_details = extract_proposed_trademark_details2(temp_file_path)  
                
                if proposed_trademark_details:  
                    proposed_name = proposed_trademark_details.get('proposed_trademark_name', 'N')  
                    proposed_class = proposed_trademark_details.get('proposed_nice_classes_number')  
                    proposed_goods_services = proposed_trademark_details.get('proposed_goods_services', 'N')  
                    if (proposed_goods_services != 'N'): 
                        with st.expander(f"Proposed Trademark Details for {uploaded_file.name}"):  
                                st.write(f"Proposed Trademark name: {proposed_name}")  
                                st.write(f"Proposed class-number: {proposed_class}")  
                                st.write(f"Proposed Goods & Services: {proposed_goods_services}") 
                        class_list = list_conversion(proposed_class)  
                    else :
                        st.write("______________________________________________________________________________________________________________________________")
                        st.write(f"Sorry, unable to generate report due to insufficient information about goods & services in the original trademark report : {uploaded_file.name}")
                        st.write("______________________________________________________________________________________________________________________________")
                        sp = False
                else :  
                    st.error(f"Unable to extract Proposed Trademark Details for {uploaded_file.name}") 
                    sp = False 
                    continue  
            
            if (sp):    
                for i in range(1,21):
                    time.sleep(0.5)
                    progress_bar.progress(i)
                    
                progress_bar.progress(25)
                
                print("+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++CALLED++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++")
                existing_trademarks = parse_trademark_details(temp_file_path)
                
                for i in range(25,46):
                    time.sleep(0.5)
                    progress_bar.progress(i)  
                    
                progress_bar.progress(50)
                st.success(f"Existing Trademarks Data Extracted Successfully for {uploaded_file.name}!")  
                # Display extracted details              
                
                nfiltered_list = []
                
                # Iterate over each JSON element in trademark_name_list  
                for json_element in existing_trademarks:  
                    class_numbers = json_element["international_class_number"]  
                # Check if any of the class numbers are in class_list  
                    if any(number in class_list for number in class_numbers):  
                        nfiltered_list.append(json_element)
                    
                existing_trademarks = nfiltered_list
                     
                high_conflicts = []
                moderate_conflicts = []
                low_conflicts = []
                
                lt = len(existing_trademarks)
                
                for existing_trademark in existing_trademarks:  
                    conflict = compare_trademarks(existing_trademark, proposed_name, proposed_class, proposed_goods_services)  
                    if conflict['conflict_grade'] == "High":  
                        high_conflicts.append(conflict)  
                    elif conflict['conflict_grade'] == "Moderate":  
                        moderate_conflicts.append(conflict)  
                    else:  
                        low_conflicts.append(conflict)  
    
                st.sidebar.write("_________________________________________________")
                st.sidebar.subheader("\n\nConflict Grades : \n")  
                st.sidebar.markdown(f"File: {proposed_name}")  
                st.sidebar.markdown(f"Total number of conflicts: {len(high_conflicts) + len(moderate_conflicts) + len(low_conflicts)}")
                st.sidebar.markdown(f"High Conflicts: {len(high_conflicts)}")  
                st.sidebar.markdown(f"Moderate Conflicts: {len(moderate_conflicts)}")  
                st.sidebar.markdown(f"Low Conflicts: {len(low_conflicts)}")  
                st.sidebar.write("_________________________________________________")
    
                document = Document()  
                
                document.add_heading(f'Trademark Conflict List for {proposed_name} :')            
                document.add_paragraph(f"\n\nTotal number of conflicts: {len(high_conflicts) + len(moderate_conflicts) + len(low_conflicts)}\n- High Conflicts: {len(high_conflicts)}\n- Moderate Conflicts: {len(moderate_conflicts)}\n- Low Conflicts: {len(low_conflicts)}\n")  
                
                if len(high_conflicts) > 0:  
                            document.add_heading('Trademarks with High Conflicts:', level=2)  
                            # Create a pandas DataFrame from the JSON list    
                            df_high = pd.DataFrame(high_conflicts) 
                            df_high = df_high.drop(columns=['Trademark serial number','reasoning'])  
                            # Create a table in the Word document    
                            table_high = document.add_table(df_high.shape[0] + 1, df_high.shape[1])
                            # Set a predefined table style (with borders)  
                            table_high.style = 'TableGrid'  # This is a built-in style that includes borders  
                            # Add the column names to the table    
                            for i, column_name in enumerate(df_high.columns):  
                                table_high.cell(0, i).text = column_name  
                            # Add the data to the table    
                            for i, row in df_high.iterrows():  
                                for j, value in enumerate(row):  
                                    table_high.cell(i + 1, j).text = str(value)

                if len(moderate_conflicts) > 0:  
                            document.add_heading('Trademarks with Moderate Conflicts:', level=2)  
                            # Create a pandas DataFrame from the JSON list    
                            df_moderate = pd.DataFrame(moderate_conflicts)
                            df_moderate = df_moderate.drop(columns=['Trademark serial number','reasoning'])  
                            # Create a table in the Word document    
                            table_moderate = document.add_table(df_moderate.shape[0] + 1, df_moderate.shape[1])
                            # Set a predefined table style (with borders)  
                            table_moderate.style = 'TableGrid'  # This is a built-in style that includes borders  
                            # Add the column names to the table    
                            for i, column_name in enumerate(df_moderate.columns):  
                                table_moderate.cell(0, i).text = column_name  
                            # Add the data to the table    
                            for i, row in df_moderate.iterrows():  
                                for j, value in enumerate(row):  
                                    table_moderate.cell(i + 1, j).text = str(value)

                if len(low_conflicts) > 0:  
                            document.add_heading('Trademarks with Low Conflicts:', level=2)  
                            # Create a pandas DataFrame from the JSON list    
                            df_low = pd.DataFrame(low_conflicts)  
                            df_low = df_low.drop(columns=['Trademark serial number', 'reasoning'])
                            # Create a table in the Word document    
                            table_low = document.add_table(df_low.shape[0] + 1, df_low.shape[1])
                            # Set a predefined table style (with borders)  
                            table_low.style = 'TableGrid'  # This is a built-in style that includes borders  
                            # Add the column names to the table    
                            for i, column_name in enumerate(df_low.columns):  
                                table_low.cell(0, i).text = column_name  
                            # Add the data to the table    
                            for i, row in df_low.iterrows():  
                                for j, value in enumerate(row):  
                                    table_low.cell(i + 1, j).text = str(value)
                            
                def add_conflict_paragraph(document, conflict):  
                    p = document.add_paragraph(f"Trademark Name : {conflict.get('Trademark name', 'N/A')}")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(f"Trademark Status : {conflict.get('Trademark status', 'N/A')}")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(f"Trademark Owner : {conflict.get('Trademark owner', 'N/A')}")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(f"Trademark Class Number : {conflict.get('Trademark class Number', 'N/A')}")  
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)  
                    p = document.add_paragraph(f"Trademark serial number : {conflict.get('Trademark serial number', 'N/A')}")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(" ")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p.paragraph_format.space_after = Pt(0) 
                    p = document.add_paragraph(f"{conflict.get('reasoning','N/A')}\n")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p = document.add_paragraph(" ")  
                    p.paragraph_format.line_spacing = Pt(18)  
                
                if len(high_conflicts) > 0:  
                    document.add_heading('Trademarks with High Conflicts Reasoning:', level=2)  
                    p = document.add_paragraph(" ")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    for conflict in high_conflicts:  
                        add_conflict_paragraph(document, conflict)  
                
                if len(moderate_conflicts) > 0:  
                    document.add_heading('Trademarks with Moderate Conflicts Reasoning:', level=2)  
                    p = document.add_paragraph(" ")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    for conflict in moderate_conflicts:  
                        add_conflict_paragraph(document, conflict)  
                
                if len(low_conflicts) > 0:  
                    document.add_heading('Trademarks with Low Conflicts Reasoning:', level=2)  
                    p = document.add_paragraph(" ")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    for conflict in low_conflicts:  
                        add_conflict_paragraph(document, conflict)  
                        
                for i in range(70,96):
                    time.sleep(0.5)
                    progress_bar.progress(i)  
                    
                progress_bar.progress(100)
    
                filename = proposed_name
                doc_stream = BytesIO()  
                document.save(doc_stream)  
                doc_stream.seek(0)  
                download_table = f'<a href="data:application/octet-stream;base64,{base64.b64encode(doc_stream.read()).decode()}" download="{filename + " Trademark Conflict Report"}.docx">Download: {filename}</a>'  
                st.sidebar.markdown(download_table, unsafe_allow_html=True)  
                st.success(f"{proposed_name} Document conflict report successfully completed!")
                st.write("______________________________________________________________________________________________________________________________")
  
        progress_bar.progress(100)
        st.success("All documents processed successfully!")  
